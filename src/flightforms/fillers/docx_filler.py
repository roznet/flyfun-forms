"""DOCX filler for LFQA (Reims Prunay) customs form.

Template structure:
- Table 0: arrival/departure checkbox (cols 1 and 4)
- Table 1: flight details (col 0 = arrival side, col 2 = departure side)
- Table 2: crew (rows 2-4 = data rows, 5 cols)
- Table 3: passengers (rows 2-4 = data rows, 5 cols)
- P14: observations
"""

from copy import deepcopy
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.table import _Row

from ..api.models import GenerateRequest
from ..registry import FormMapping


def _parse_date(date_str: str, fmt: str) -> str:
    dt = datetime.strptime(date_str, "%Y-%m-%d")
    return dt.strftime(fmt)


def _append_after_colon(paragraph, value: str):
    """Append value after the last colon in the paragraph, preserving formatting."""
    text = paragraph.text
    if ":" in text:
        # Add a new run with the value after the existing text
        run = paragraph.add_run(f" {value}")
        run.font.size = paragraph.runs[0].font.size if paragraph.runs else None


def _fill_flight_details_cell(cell, data: dict, mapping: FormMapping):
    """Fill the flight details in a table 1 cell (arrival or departure side)."""
    for paragraph in cell.paragraphs:
        text = paragraph.text
        if "Date" in text and "Heure" in text:
            _append_after_colon(paragraph, f"{data['date']}  {data['time']}")
        elif "Propriétaire" in text or "Owner" in text:
            parts = []
            if data.get("owner"):
                parts.append(data["owner"])
            parts.append(data["registration"])
            parts.append(data["aircraft_type"])
            _append_after_colon(paragraph, "  ".join(parts))
        elif "Provenance" in text or "Destination" in text:
            airport_text = f"{data['airport']} ({data['airport_name']})"
            _append_after_colon(paragraph, airport_text)
        elif "Pays" in text:
            _append_after_colon(paragraph, data.get("country", ""))
        elif "Nature" in text:
            _append_after_colon(paragraph, data.get("nature", ""))


def _add_table_row(table) -> _Row:
    """Add a new row to a table, copying the structure of the last row."""
    last_row = table.rows[-1]
    new_row = deepcopy(last_row._tr)
    table._tbl.append(new_row)
    return table.rows[-1]


def fill_docx(
    template_path: Path,
    mapping: FormMapping,
    request: GenerateRequest,
    airport_resolver,
) -> bytes:
    doc = Document(str(template_path))
    date_fmt = mapping.date_format
    is_arrival = request.airport == request.flight.destination
    observations = request.observations or mapping.default_observations or ""

    # Table 0: arrival/departure mark
    table0 = doc.tables[0]
    if is_arrival:
        table0.rows[0].cells[1].text = "X"
    else:
        table0.rows[0].cells[4].text = "X"

    # Table 1: flight details
    table1 = doc.tables[1]

    # Arrival side (col 0)
    arrival_data = {
        "date": _parse_date(request.flight.arrival_date, date_fmt),
        "time": request.flight.arrival_time_utc,
        "owner": request.aircraft.owner or "",
        "registration": request.aircraft.registration,
        "aircraft_type": request.aircraft.type,
        "airport": request.flight.origin,
        "airport_name": airport_resolver.get_name(request.flight.origin),
        "country": airport_resolver.get_country(request.flight.origin),
        "nature": request.flight.nature,
    }
    _fill_flight_details_cell(table1.rows[0].cells[0], arrival_data, mapping)

    # Departure side (col 2)
    departure_data = {
        "date": _parse_date(request.flight.departure_date, date_fmt),
        "time": request.flight.departure_time_utc,
        "owner": request.aircraft.owner or "",
        "registration": request.aircraft.registration,
        "aircraft_type": request.aircraft.type,
        "airport": request.flight.destination,
        "airport_name": airport_resolver.get_name(request.flight.destination),
        "country": airport_resolver.get_country(request.flight.destination),
        "nature": request.flight.nature,
    }
    _fill_flight_details_cell(table1.rows[0].cells[2], departure_data, mapping)

    # Table 2: crew
    _fill_person_table(doc.tables[2], request.crew, date_fmt)

    # Table 3: passengers
    _fill_person_table(doc.tables[3], request.passengers, date_fmt)

    # Observations
    for para in doc.paragraphs:
        if para.text.strip().startswith("OBSERVATIONS"):
            para.add_run(f" {observations}")
            break

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _fill_person_table(table, people: list, date_fmt: str):
    """Fill crew or passenger table. Rows 2+ are data rows (row 0=header, row 1=column labels)."""
    data_start = 2  # First data row index

    # Ensure enough rows
    existing_data_rows = len(table.rows) - data_start
    for _ in range(len(people) - existing_data_rows):
        _add_table_row(table)

    for i, person in enumerate(people):
        row = table.rows[data_start + i]
        row.cells[0].text = person.last_name
        row.cells[1].text = person.first_name
        row.cells[2].text = _parse_date(person.dob, date_fmt) if person.dob else ""
        row.cells[3].text = person.nationality or ""
        row.cells[4].text = person.id_number or ""

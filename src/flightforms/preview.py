"""Preview generation with self-describing dummy data and field extraction.

Used by the CLI ``preview`` command and by snapshot regression tests.
Each dummy value encodes its canonical field name so visual inspection
immediately reveals whether a value landed in the right place.
"""

from datetime import datetime
from io import BytesIO
from pathlib import Path

from .api.models import (
    AircraftData,
    ConnectingFlightData,
    FlightData,
    GenerateRequest,
    PersonData,
)
from .registry import FormMapping, MappingRegistry


# ── Stub resolver for preview (no DB needed) ────────────────────────────────

class PreviewAirportResolver:
    """Returns recognisable names without needing the euro_aip database."""

    _NAMES = {
        "ORIG": "OriginCity",
        "DEST": "DestCity",
        "CNXN": "ConnectCity",
    }
    _COUNTRIES = {
        "ORIG": "OriginCountry",
        "DEST": "DestCountry",
        "CNXN": "ConnectCountry",
    }

    def get_name(self, icao: str) -> str:
        return self._NAMES.get(icao, f"{icao}City")

    def get_country(self, icao: str) -> str:
        return self._COUNTRIES.get(icao, f"{icao}Country")

    def get_country_code(self, icao: str) -> str:
        return ""


# ── Dummy data builders ─────────────────────────────────────────────────────

def _make_person(prefix: str, index: int) -> PersonData:
    """Build a person whose values encode the field name + index."""
    n = index + 1
    return PersonData(
        function=f"{prefix}Func{n}",
        first_name=f"{prefix}First{n}",
        last_name=f"{prefix}Last{n}",
        dob=f"200{n}-0{n}-1{n}",
        nationality=f"{prefix}Nat{n}",
        id_number=f"{prefix}Id{n}",
        id_type=f"{prefix}IdType{n}",
        id_issuing_country=f"{prefix}IdCountry{n}",
        id_expiry=f"203{n}-0{n}-1{n}",
        sex="Male" if n % 2 else "Female",
        place_of_birth=f"{prefix}POB{n}",
        address=f"{prefix}Addr{n}",
    )


def make_preview_request(
    mapping: FormMapping,
    airport: str,
    direction: str = "arrival",
) -> GenerateRequest:
    """Build a request with self-describing dummy values for *mapping*.

    *direction* is ``"arrival"`` (flight into *airport*) or ``"departure"``
    (flight out of *airport*).
    """
    if direction == "arrival":
        origin = "ORIG"
        destination = airport
    else:
        origin = airport
        destination = "DEST"

    num_crew = min(mapping.max_crew, 2)
    num_pax = min(mapping.max_passengers, 2)

    crew = [_make_person("Crew", i) for i in range(num_crew)]
    pax = [_make_person("Pax", i) for i in range(num_pax)]

    extra_fields = None
    if mapping.extra_fields:
        extra_fields = {}
        for ef in mapping.extra_fields:
            ef_def = ef if isinstance(ef, dict) else {"key": ef, "type": "text"}
            key = ef_def["key"]
            ef_type = ef_def.get("type", "text")
            if ef_type == "choice":
                options = ef_def.get("options", [])
                extra_fields[key] = options[0] if options else f"Extra{key}"
            elif ef_type == "person":
                extra_fields[key] = {
                    "name": f"Extra{key}Name",
                    "address": f"Extra{key}Addr",
                }
            else:
                extra_fields[key] = f"Extra{key}"

    connecting = None
    if mapping.has_connecting_flight:
        connecting = ConnectingFlightData(
            origin=airport,
            destination="CNXN",
            departure_date="2099-02-20",
            departure_time_utc="14:00",
            arrival_date="2099-02-20",
            arrival_time_utc="16:30",
        )

    return GenerateRequest(
        airport=airport,
        form=mapping.id,
        flight=FlightData(
            origin=origin,
            destination=destination,
            departure_date="2099-01-15",
            departure_time_utc="08:30",
            arrival_date="2099-01-15",
            arrival_time_utc="10:45",
            nature="private",
            contact="PreviewContact",
        ),
        aircraft=AircraftData(
            registration="AcReg",
            type="AcType",
            owner="AcOwner",
            owner_address="AcOwnerAddr",
            is_airplane=True,
            usual_base="ORIG",
        ),
        crew=crew,
        passengers=pax,
        extra_fields=extra_fields,
        connecting_flight=connecting,
        observations="PreviewObs",
    )


# ── Form generation ─────────────────────────────────────────────────────────

# Airport to use when generating each form (must satisfy direction constraints)
def _find_mapping_by_id(registry: MappingRegistry, form_id: str) -> FormMapping | None:
    """Search all loaded mappings by ID, falling back to loading directly from JSON.

    Some forms (e.g. gendec_icao) are manually-selectable and not auto-discovered
    by the registry, so we load them from the mappings directory as a fallback.
    """
    for mappings in registry.all_airports().values():
        for m in mappings:
            if m.id == form_id:
                return m
    for mappings in registry.all_prefixes().values():
        for m in mappings:
            if m.id == form_id:
                return m
    for m in registry.all_defaults():
        if m.id == form_id:
            return m
    # Direct file load for unscoped mappings
    import json
    json_path = registry.mappings_dir / f"{form_id}.json"
    if json_path.exists():
        with open(json_path) as f:
            data = json.load(f)
        return FormMapping(data, form_id)
    return None


FORM_AIRPORTS = {
    "lsgs": "LSGS",
    "french_customs": "LFAC",
    "lfqa": "LFQA",
    "gar": "EGKA",
    "gendec_form": "DEST",
    "gendec_icao": "DEST",
}

# Forms where arrival vs departure changes what appears on the form
DIRECTION_AWARE_FORMS = {"lsgs", "french_customs", "lfqa", "gar"}


def generate_preview(
    registry: MappingRegistry,
    form_id: str,
    resolver=None,
    direction: str = "arrival",
) -> bytes:
    """Generate a filled form using self-describing dummy data.

    *direction* is ``"arrival"`` or ``"departure"``.
    Returns the raw document bytes (PDF, DOCX, or XLSX).
    """
    from .fillers.pdf_filler import fill_pdf
    from .fillers.french_customs_filler import fill_french_customs
    from .fillers.docx_filler import fill_docx
    from .fillers.xlsx_filler import fill_xlsx

    if resolver is None:
        resolver = PreviewAirportResolver()

    airport = FORM_AIRPORTS.get(form_id, "ZZZZ")
    mapping = registry.get_form(airport, form_id)
    if mapping is None:
        mapping = _find_mapping_by_id(registry, form_id)
    if mapping is None:
        raise ValueError(f"Form {form_id!r} not found")

    request = make_preview_request(mapping, airport, direction=direction)
    template_path = registry.get_template_path(mapping)

    fillers = {
        "pdf_acroform": lambda: fill_pdf(template_path, mapping, request, resolver),
        "pdf_acroform_french": lambda: fill_french_customs(template_path, mapping, request, resolver),
        "docx": lambda: fill_docx(template_path, mapping, request, resolver),
        "xlsx": lambda: fill_xlsx(template_path, mapping, request, resolver),
    }

    filler = fillers.get(mapping.filler_type)
    if filler is None:
        raise ValueError(f"Unknown filler type {mapping.filler_type!r}")
    return filler()


# ── Field extraction utilities ───────────────────────────────────────────────

def extract_pdf_fields(pdf_bytes: bytes) -> dict[str, str]:
    """Extract AcroForm field name → value from a filled PDF."""
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(pdf_bytes))
    fields = reader.get_fields() or {}
    result = {}
    for name, field in fields.items():
        value = field.get("/V")
        if value is not None:
            result[name] = str(value)
    return result


def extract_xlsx_fields(xlsx_bytes: bytes, sheet: str = "GAR") -> dict[str, str]:
    """Extract cell address → value for every non-empty cell in *sheet*."""
    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(xlsx_bytes), data_only=True)
    ws = wb[sheet]
    result = {}
    for row in ws.iter_rows():
        for cell in row:
            if cell.value is not None:
                val = cell.value
                if isinstance(val, datetime):
                    val = val.strftime("%Y-%m-%d")
                result[cell.coordinate] = str(val)
    return result


def extract_docx_tables(docx_bytes: bytes) -> list[list[list[str]]]:
    """Extract table[i] → row[j] → cell[k] text from a DOCX.

    Returns a nested list: tables[table_idx][row_idx][cell_idx] = text.
    """
    from docx import Document

    doc = Document(BytesIO(docx_bytes))
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            rows.append(cells)
        tables.append(rows)
    return tables


def extract_docx_paragraphs(docx_bytes: bytes) -> list[str]:
    """Extract all paragraph texts from a DOCX."""
    from docx import Document

    doc = Document(BytesIO(docx_bytes))
    return [p.text for p in doc.paragraphs]


# ── File extension helper ────────────────────────────────────────────────────

FILLER_EXTENSIONS = {
    "pdf_acroform": ".pdf",
    "pdf_acroform_french": ".pdf",
    "docx": ".docx",
    "xlsx": ".xlsx",
}
